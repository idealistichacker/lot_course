# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Title
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('物联网与边缘计算（报告）\n基于边缘RAG的海滨阳台气象与资产哨兵系统')
    run.font.size = Pt(16)
    run.font.bold = True

    # Info
    info = doc.add_paragraph()
    info.add_run('成    绩：              \n')
    info.add_run('学生姓名：              \n')
    info.add_run('学    号：              \n')
    info.add_run('专业班级：              \n')
    info.add_run('同组者姓名：                  \n\n')
    info.add_run('2026年 5月20日\n')

    # Abstract
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('摘  要')
    run.font.size = Pt(14)
    run.font.bold = True

    p = doc.add_paragraph('本项目围绕“海滨阳台无人值守气象与资产哨兵”展开，目标是在宿舍无人照看时，让终端设备自动采集阳台环境和摄像头状态，再由树莓派完成本地知识库检索、推理判断和移动端告警。系统前端感知节点采用 ESP32-S3 CAM 与 DHT11，负责温湿度、露点、Wi-Fi 状态和摄像头初始化信息的上报；树莓派作为边缘计算节点，接收设备数据，结合本地规则、课程节律和资产脆弱性知识，生成决策。本报告重点说明本人承担的 ESP32 数据采集与上传部分，尤其是底层传感器的驱动、网络连接的稳定性维护、数据封包上报机制以及深度休眠功耗管理。')
    p = doc.add_paragraph('关键词：物联网；ESP32-S3 CAM；数据采集；HTTP 上传；深度休眠')

    # TOC
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('目  录')
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph('摘  要')
    doc.add_paragraph('第1章  引言')
    doc.add_paragraph('第2章  总体设计')
    doc.add_paragraph('第3章  ESP32终端软硬件初始化')
    doc.add_paragraph('第4章  数据采集与露点计算')
    doc.add_paragraph('第5章  网络通信与可靠上传')
    doc.add_paragraph('第6章  功耗管理与休眠唤醒')
    doc.add_paragraph('第7章  心得体会')

    # Abstract formatting
    for p in doc.paragraphs:
        p.paragraph_format.first_line_indent = Inches(0.3)

    # Ch 1
    doc.add_page_break()
    h = doc.add_paragraph()
    run = h.add_run('第1章  引言')
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph('海边宿舍的阳台环境变化比较快，风、湿度、盐雾和短时降雨都可能影响放在外面的物品。本项目旨在构建一个端到端的无人值守哨兵系统，通过边缘计算与物联网技术，实现阳台环境的实时感知与智能告警。在此链路中，稳定的前端数据采集是后续边缘推理与决策的基石。')
    doc.add_paragraph('本次项目中，我的工作主要集中在感知层，即负责基于 ESP32-S3 CAM 的端侧开发。这不仅包含 DHT11 温湿度数据的读取，还涉及关键指标（如露点）的边缘预计算、摄像头的适配与初始化、Wi-Fi 网络的建立及断线重连保障、HTTP 协议的数据封包发送，以及整机的低功耗休眠调度。队友则主要负责树莓派侧的数据接收、知识库匹配和基于大模型的最终决策。我负责的 ESP32 采集上传部分与队友的树莓派边缘计算部分通过标准化的 JSON 报文紧密衔接，共同构成了项目的完整闭环。')

    # Ch 2
    h = doc.add_paragraph()
    run = h.add_run('第2章  总体设计')
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph('系统架构自下而上划分为感知层、边缘层。感知层由我负责，以 ESP32-S3 CAM 和 DHT11 为核心，肩负着现场“哨兵”的最基础任务：在固定节律下唤醒，捕获环境物理量及设备自检状态，将其组装成语义明确的报文后推送到局域网内。')
    doc.add_paragraph('从设计哲学看，终端设备不必承担繁重的推理运算，但必须对自身的通信可靠性及工作状态有清晰的记录（如启动次数、序列号、在线时长等）。数据流出后送入边缘层，交由树莓派处理。对于终端来说，主要工作流被严格限定为“唤醒 -> 初始化网络及传感器 -> 尝试采集数据及拍照自检 -> 组装并发送报文 -> 进入深度休眠”，这种无状态、触发式的单次运行模式极大提高了弱网环境下的系统存活率和自我恢复能力。')

    # Ch 3
    h = doc.add_paragraph()
    run = h.add_run('第3章  ESP32终端软硬件初始化')
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph('在硬件选型层面，ESP32-S3 CAM 具备强大的计算能力并自带了能够评估阳台现场情况的图像接口。在软开机阶段，首要任务是保障系统时钟与各类芯片控制器的正常拉起。每次从深度休眠中返回时，程序除了对串口实施标准波特率（115200）设定外，还会通过 RTC 变量来持久化核心计数器。比如 `bootCount`（启动次数）和 `sequenceId`（运行序列号）都声明为 `RTC_DATA_ATTR`。这样，每次冷启动或休眠唤醒时，这两个变量不会丢失。它们被一同装载进后续要发出的数据包中，能直观反映出模块是否有意外宕机或是正在平稳地周期运转。')
    doc.add_paragraph('摄像头的初始化相对复杂，牵涉到引脚矩阵的精确映射。配置时需要映射包括 XCLK、PCLK、VSYNC 在内的一系列时序及数据引脚，并指定频率等参数。若底层存在 PSRAM，程序将自动以较高分辨率（FRAMESIZE_VGA）启动并开辟双缓存；反之则降级运行以保障系统不至于内存溢出崩溃。由于目前主要联调焦点在环境数据与规则引擎间，图像资源的上传虽暂缓接入，但在硬件与初始化阶段我已提前确保了摄像头驱动层的链路连通性，`camera_ok` 的状态指示已正式纳入传输协议，为后期的全功能开放提供了底层支撑。')

    # Ch 4
    h = doc.add_paragraph()
    run = h.add_run('第4章  数据采集与露点计算')
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph('在核心环境参数采集环节，我利用 DHT11 传感器进行温度与湿度的读取。阳台环境恶劣且可能遭遇电磁干扰，传感器偶有超时或奇偶校验不通过导致读出无效值 (NaN) 的情况。为屏蔽这种硬件偶发的不稳定性，代码中没有采用一读定音的方式，而是专门封装了一层带冗余重试逻辑的读取机制。在至多5次循环尝试中，设定 1500 毫秒的延时间隙以等待 DHT11 从上一次应答中恢复，只有读出合法数值且湿度在 0~100 之间，才向主程序层通报采集成功 (Sensor OK)。')
    doc.add_paragraph('此外考虑到海滨阳台上物品的“受损机制”，通常不仅取决于绝对湿度，更容易因“露点降温”引发结露。因此，我在前端承担了轻量级的局部推算。引入了一个工程近似公式，利用读取到的当前摄氏温度减去 (100 - 相对湿度百分比) 除以 5 的差值，求出近似露点温度。提前在 ESP32 上计算好本批次数据的露点，不但保证了上报结构体的完整性，还变相减轻了边缘侧或大模型进行直接物理量折算的压力。')

    # Ch 5
    h = doc.add_paragraph()
    run = h.add_run('第5章  网络通信与可靠上传')
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph('对于物联网节点而言，不发往服务端的数据毫无意义。网络连接部分的鲁棒性是这部分的重中之重。ESP32 被配置为专职的 STA 模式，为摸排周遭无线信道状况并在日后排障中提供依据，在建立连接前会执行一次周边网络扫描，并在串行端口输出捕获列表的详细信息。')
    doc.add_paragraph('随后系统依靠预编译的 SSID 与密码发起接入，并配有长达20秒的看门等待。当网络状态成功置为 `WL_CONNECTED`，说明路由链路就绪；失败则记录状态码告警。为了将各维度数据结构化，我手写了一套高效的 JSON 组装逻辑，将设备 ID、固件版本号与刚刚采得的时间、指标数据打包。')
    doc.add_paragraph('数据发送阶段同样采取带有重试补偿特性的 `HTTPClient` 调用机制，甚至在失败回调里兼顾了 Wi-Fi 骤然掉线的重连尝试。在最多 3 次 POST 尝试内，只要得到 200 族 HTTP 应答码即视为投递完成，确保不会因一次服务器丢包而抛弃整个周期内的环境现场信息。最终发出的报文不仅能对接现有的联调 Mock 端，更遵循了与队友约定的格式规范，为树莓派侧知识库提供了充足的特征依据。')

    # Ch 6
    h = doc.add_paragraph()
    run = h.add_run('第6章  功耗管理与休眠唤醒')
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph('阳台哨兵受限于供电苛刻的实际情况，持续全负载运行显然是不现实的。在流程的收尾阶段，不管 HTTP 上传是否成功，都会进行资源释放步骤。这包括强制切断当前 Wi-Fi 的底层维持任务并将无线模式关闭，切断大部分外设电路以节省毫安级别的消耗。')
    doc.add_paragraph('接着系统会载入下一个唤醒闹钟并启动 Deep Sleep，挂起绝大多数片上时钟系统与应用层线程。这种基于定时器强关联唤醒的“脉冲式”工作特征，使哨兵可以每过一段特定的宏观时间（目前约定约为 15 分钟）重启轮回，极大拉伸了可能由锂电池或蓄电板接驳供电系统时的工作生命周期。这也是此类被动部署物联网终端有别于常驻路由与核心网关最大的系统特征之一。')

    # Ch 7
    h = doc.add_paragraph()
    run = h.add_run('第7章  心得体会')
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph('在这个项目中，我深刻意识到端侧系统的开发并不仅仅是跑通 Demo 那么简单，尤其是在真实的外部环境长期部署时，程序必须对各种时序异常、传感器卡死以及网络质量波动有极强的容忍底线与自限机制。通过把“重试”、“验证有效性”与“硬件看门狗意识”融入代码的字里行间，我获得了在极端情况下保持数据活性的宝贵经验。')
    doc.add_paragraph('同时我也深刻感受到与队友的上下游配合是何等关键。我传出的不再是几个割裂的浮点数，而是承载了一方天地的数字镜像和状态上下文；通过提供类似 WiFi RSSI 以及露点等深挖的环境物理量，令后置 RAG 系统在判断时拥有了更多颗粒度。看到从我这边封装起的一个 JSON 顺利在树莓派中唤起分析进程到最终生成业务判断时，我切实尝到了软硬件系统集成的成就感，这也为我未来探究更高阶的视觉边缘上报与嵌入式架构优化指明了航向。')

    doc.save('物联网与边缘计算报告_esp32数据采集上传部分.docx')

if __name__ == '__main__':
    create_report()
